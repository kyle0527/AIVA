#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AIVA Word 文檔轉 Markdown 工具
將 .docx 檔案轉換為 .md 格式並進行內容分析
"""

import sys
import os
from pathlib import Path
from docx import Document
import re
from typing import List, Dict, Any

class DocxToMarkdownConverter:
    """Word 文檔轉 Markdown 轉換器"""
    
    def __init__(self):
        self.content_stats = {
            'total_paragraphs': 0,
            'total_tables': 0,
            'total_images': 0,
            'headings': {},
            'key_sections': []
        }
    
    def convert_docx_to_markdown(self, docx_path: str, output_path: str = None) -> str:
        """
        將 Word 文檔轉換為 Markdown
        
        Args:
            docx_path: Word 文檔路徑
            output_path: 輸出 Markdown 檔案路徑
        
        Returns:
            Markdown 內容字串
        """
        try:
            # 讀取 Word 文檔
            doc = Document(docx_path)
            markdown_content = []
            
            # 處理文檔標題
            if hasattr(doc, 'core_properties') and doc.core_properties.title:
                markdown_content.append(f"# {doc.core_properties.title}\n")
            
            # 處理段落
            for paragraph in doc.paragraphs:
                self.content_stats['total_paragraphs'] += 1
                md_line = self._convert_paragraph_to_markdown(paragraph)
                if md_line.strip():
                    markdown_content.append(md_line)
            
            # 處理表格
            for table in doc.tables:
                self.content_stats['total_tables'] += 1
                md_table = self._convert_table_to_markdown(table)
                markdown_content.append(md_table)
            
            # 合併內容
            full_markdown = '\n'.join(markdown_content)
            
            # 如果指定了輸出路徑，寫入檔案
            if output_path:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(full_markdown)
                print(f"✅ Markdown 檔案已保存至: {output_path}")
            
            return full_markdown
            
        except Exception as e:
            print(f"❌ 轉換失敗: {e}")
            return ""
    
    def _convert_paragraph_to_markdown(self, paragraph) -> str:
        """將段落轉換為 Markdown 格式"""
        if not paragraph.text.strip():
            return ""
        
        text = paragraph.text.strip()
        
        # 檢查是否為標題
        if paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
            self.content_stats['headings'][level] = self.content_stats['headings'].get(level, 0) + 1
            return f"{'#' * level} {text}\n"
        
        # 檢查是否為列表項
        if paragraph.style.name in ['List Paragraph', 'ListParagraph']:
            return f"- {text}\n"
        
        # 一般段落
        return f"{text}\n"
    
    def _convert_table_to_markdown(self, table) -> str:
        """將表格轉換為 Markdown 格式"""
        if not table.rows:
            return ""
        
        markdown_table = []
        
        # 處理表頭
        header_row = table.rows[0]
        headers = [cell.text.strip() for cell in header_row.cells]
        markdown_table.append("| " + " | ".join(headers) + " |")
        markdown_table.append("| " + " | ".join(["---"] * len(headers)) + " |")
        
        # 處理資料行
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            markdown_table.append("| " + " | ".join(cells) + " |")
        
        return "\n".join(markdown_table) + "\n\n"
    
    def analyze_content(self, markdown_content: str) -> Dict[str, Any]:
        """分析 Markdown 內容"""
        analysis = {
            'document_stats': self.content_stats.copy(),
            'content_analysis': {},
            'key_topics': [],
            'technical_sections': [],
            'recommendations': []
        }
        
        # 字數統計
        word_count = len(markdown_content.split())
        char_count = len(markdown_content)
        line_count = len(markdown_content.split('\n'))
        
        analysis['content_analysis'] = {
            'word_count': word_count,
            'character_count': char_count,
            'line_count': line_count,
            'estimated_reading_time': f"{word_count // 200} 分鐘"
        }
        
        # 尋找關鍵主題
        key_patterns = [
            r'多語言.*?程式',
            r'AIVA.*?系統',
            r'Python.*?Go.*?Rust',
            r'安全.*?平台',
            r'架構.*?設計',
            r'性能.*?優化',
            r'模組.*?整合'
        ]
        
        for pattern in key_patterns:
            matches = re.findall(pattern, markdown_content, re.IGNORECASE)
            if matches:
                analysis['key_topics'].extend(matches[:3])  # 最多取 3 個匹配
        
        # 尋找技術段落
        tech_patterns = [
            r'.*?API.*?',
            r'.*?架構.*?',
            r'.*?模組.*?',
            r'.*?系統.*?',
            r'.*?實現.*?',
            r'.*?優化.*?'
        ]
        
        lines = markdown_content.split('\n')
        for line in lines:
            if any(re.search(pattern, line, re.IGNORECASE) for pattern in tech_patterns):
                if len(line.strip()) > 10:  # 過濾太短的行
                    analysis['technical_sections'].append(line.strip())
        
        # 限制技術段落數量
        analysis['technical_sections'] = analysis['technical_sections'][:10]
        
        return analysis
    
    def generate_analysis_report(self, analysis: Dict[str, Any], markdown_content: str) -> str:
        """產生分析報告"""
        report = []
        
        report.append("# 📊 AIVA 多語言程式處理能力實現研究 - 分析報告\n")
        
        # 文檔統計
        report.append("## 📈 文檔統計資訊")
        stats = analysis['content_analysis']
        report.append(f"- **字數**: {stats['word_count']:,} 字")
        report.append(f"- **字元數**: {stats['character_count']:,} 字元")
        report.append(f"- **行數**: {stats['line_count']:,} 行")
        report.append(f"- **預估閱讀時間**: {stats['estimated_reading_time']}")
        report.append(f"- **段落數**: {analysis['document_stats']['total_paragraphs']}")
        report.append(f"- **表格數**: {analysis['document_stats']['total_tables']}")
        report.append("")
        
        # 標題結構
        if analysis['document_stats']['headings']:
            report.append("## 📋 文檔結構")
            for level, count in sorted(analysis['document_stats']['headings'].items()):
                report.append(f"- 第 {level} 級標題: {count} 個")
            report.append("")
        
        # 關鍵主題
        if analysis['key_topics']:
            report.append("## 🎯 關鍵主題")
            for topic in set(analysis['key_topics'][:5]):  # 去重並限制數量
                report.append(f"- {topic}")
            report.append("")
        
        # 技術重點
        if analysis['technical_sections']:
            report.append("## 🔧 技術重點段落")
            for i, section in enumerate(analysis['technical_sections'][:5], 1):
                report.append(f"{i}. {section}")
            report.append("")
        
        # 內容預覽
        report.append("## 📄 文檔內容預覽")
        lines = markdown_content.split('\n')
        preview_lines = [line for line in lines[:20] if line.strip()][:10]
        for line in preview_lines:
            if line.startswith('#'):
                report.append(f"\n{line}")
            else:
                report.append(f"   {line}")
        
        if len(lines) > 20:
            report.append("\n   ... (內容已截斷)")
        
        report.append("")
        
        return '\n'.join(report)

def main():
    """主函數"""
    # 設定檔案路徑
    docx_file = r"C:\D\fold7\AIVA-git\_out\AIVA 多語言程式處理能力的實現研究.docx"
    md_file = r"C:\D\fold7\AIVA-git\_out\AIVA_多語言程式處理能力的實現研究.md"
    analysis_file = r"C:\D\fold7\AIVA-git\_out\AIVA_多語言程式處理能力分析報告.md"
    
    if not os.path.exists(docx_file):
        print(f"❌ 找不到檔案: {docx_file}")
        return
    
    print(f"🔄 開始轉換 Word 文檔...")
    print(f"📁 來源檔案: {docx_file}")
    
    # 建立轉換器
    converter = DocxToMarkdownConverter()
    
    # 轉換為 Markdown
    markdown_content = converter.convert_docx_to_markdown(docx_file, md_file)
    
    if markdown_content:
        print(f"✅ 轉換完成!")
        
        # 分析內容
        print(f"🔍 正在分析內容...")
        analysis = converter.analyze_content(markdown_content)
        
        # 產生分析報告
        report = converter.generate_analysis_report(analysis, markdown_content)
        
        # 保存分析報告
        with open(analysis_file, 'w', encoding='utf-8') as f:
            f.write(report)
        
        print(f"📊 分析報告已保存至: {analysis_file}")
        
        # 顯示簡要統計
        print(f"\n📈 快速統計:")
        print(f"   - 字數: {analysis['content_analysis']['word_count']:,}")
        print(f"   - 段落: {analysis['document_stats']['total_paragraphs']}")
        print(f"   - 表格: {analysis['document_stats']['total_tables']}")
        print(f"   - 預估閱讀時間: {analysis['content_analysis']['estimated_reading_time']}")
        
    else:
        print("❌ 轉換失敗")

if __name__ == "__main__":
    main()